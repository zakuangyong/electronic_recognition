from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\project\electronic_recognition")
RESULT = ROOT / "result" / "20260621-002549-066d7b47-A17387_1706_项目原理图_10"
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "report_assets"
OUT_FILE = OUT_DIR / "电气原理图智能识别阶段成果及检索规划汇报.docx"

BLUE = "2E74B5"
NAVY = "17365D"
DARK = "1F2937"
MID = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF6F1"
GREEN = "14866D"
PALE_ORANGE = "FFF4E5"
ORANGE = "B54708"
WHITE = "FFFFFF"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, size=None, bold=None, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def style_table(table, widths=None, header=True, font_size=9):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    set_cell_width(row.cells[idx], width)
    for r_idx, row in enumerate(table.rows):
        if header and r_idx == 0:
            set_repeat_table_header(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": "D0D5DD"},
                bottom={"val": "single", "sz": 4, "color": "D0D5DD"},
                left={"val": "single", "sz": 4, "color": "D0D5DD"},
                right={"val": "single", "sz": 4, "color": "D0D5DD"},
            )
            if header and r_idx == 0:
                set_cell_fill(cell, LIGHT)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=font_size, bold=(header and r_idx == 0), color=DARK)


def add_page_title(doc, kicker, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_text(p, kicker.upper(), size=9, bold=True, color=BLUE)
    h = doc.add_heading(title, level=1)
    h.paragraph_format.space_before = Pt(0)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(10)
        add_text(p2, subtitle, size=10, color=MID)


def add_callout(doc, title, body, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_cell_width(table.cell(0, 0), 9360)
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    set_cell_border(
        cell,
        left={"val": "single", "sz": 18, "color": accent},
        top={"val": "nil"},
        bottom={"val": "nil"},
        right={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, size=11, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, body, size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(4)
        add_text(p, item, size=10, color=DARK)


def add_metric_cards(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = int(9360 / len(metrics))
    for idx, (value, label, note) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_width(cell, width)
        set_cell_fill(cell, PALE_BLUE if idx % 2 == 0 else LIGHT)
        set_cell_margins(cell, top=160, bottom=140, start=100, end=100)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 4, "color": "D0D5DD"},
            bottom={"val": "single", "sz": 4, "color": "D0D5DD"},
            left={"val": "single", "sz": 4, "color": "D0D5DD"},
            right={"val": "single", "sz": 4, "color": "D0D5DD"},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_text(p, value, size=19, bold=True, color=NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(1)
        add_text(p2, label, size=9, bold=True, color=DARK)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        add_text(p3, note, size=7.5, color=MID)


def add_image(doc, path, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    add_text(c, caption, size=8, color=MID)


def make_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    page1 = Image.open(RESULT / "pages" / "page-1.png").convert("RGB")
    page3 = Image.open(RESULT / "pages" / "page-3.png").convert("RGB")
    assets = {
        "full": (page1, ASSET_DIR / "schematic10_page1.png"),
        "main": (page1.crop((180, 120, 1250, 2200)), ASSET_DIR / "schematic10_main_circuit.png"),
        "control": (page1.crop((1220, 90, 3370, 2190)), ASSET_DIR / "schematic10_control_circuit.png"),
        "title": (page1.crop((180, 2200, 3550, 2560)), ASSET_DIR / "schematic10_title_block.png"),
        "table": (page3.crop((1950, 40, 3580, 2220)), ASSET_DIR / "schematic10_component_table.png"),
    }
    for _, (image, path) in assets.items():
        draw = ImageDraw.Draw(image)
        draw.rectangle((1, 1, image.width - 2, image.height - 2), outline=(183, 197, 214), width=3)
        image.save(path, quality=94)
    return {k: v[1] for k, v in assets.items()}


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Bullet 2"):
        st = styles[list_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10)
        st.paragraph_format.left_indent = Inches(0.25 if list_name == "List Bullet" else 0.5)
        st.paragraph_format.first_line_indent = Inches(-0.16)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    set_cell_width(table.cell(0, 0), 6500)
    set_cell_width(table.cell(0, 1), 2860)
    table.cell(0, 0).text = ""
    table.cell(0, 1).text = ""
    p = table.cell(0, 0).paragraphs[0]
    add_text(p, "电气原理图智能识别项目  |  阶段成果汇报", size=8, bold=True, color=MID)
    p2 = table.cell(0, 1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p2, "内部汇报材料", size=8, color=MID)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "2026年6月21日  |  ", size=8, color=MID)
    add_field(p, "PAGE")


def build_report():
    assets = make_assets()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "电气原理图智能识别阶段成果及检索规划汇报"
    doc.core_properties.subject = "元件识别、组合识别、文本层解析及检索规划"
    doc.core_properties.author = "电气原理图智能识别项目组"

    # Page 1 — cover / executive summary
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    add_text(p, "阶段成果汇报", size=10, bold=True, color=BLUE)
    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(4)
    add_text(title, "电气原理图智能识别项目", size=25, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    add_text(subtitle, "元件识别、组合识别、文本层解析与检索能力建设规划", size=15, color=BLUE)

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    meta_data = [
        ("汇报对象", "项目领导 / 管理团队"),
        ("汇报单位", "电气原理图智能识别项目组"),
        ("案例图纸", "原理图_10《专用排烟风机控制原理》"),
        ("汇报日期", "2026年6月21日"),
    ]
    for i, (k, v) in enumerate(meta_data):
        set_cell_width(meta.cell(i, 0), 1850)
        set_cell_width(meta.cell(i, 1), 7510)
        set_cell_fill(meta.cell(i, 0), LIGHT)
        set_cell_margins(meta.cell(i, 0), top=100, bottom=100, start=140, end=140)
        set_cell_margins(meta.cell(i, 1), top=100, bottom=100, start=140, end=140)
        set_cell_border(meta.cell(i, 0), bottom={"val": "single", "sz": 4, "color": "D0D5DD"})
        set_cell_border(meta.cell(i, 1), bottom={"val": "single", "sz": 4, "color": "D0D5DD"})
        meta.cell(i, 0).paragraphs[0].text = ""
        meta.cell(i, 1).paragraphs[0].text = ""
        add_text(meta.cell(i, 0).paragraphs[0], k, size=9, bold=True, color=MID)
        add_text(meta.cell(i, 1).paragraphs[0], v, size=10, color=DARK)

    doc.add_paragraph()
    add_callout(
        doc,
        "阶段结论",
        "目前已形成“图纸解析—开放式元件识别—知识库名称修正—组合关系判断—文本结构化输出”的基本闭环。"
        "以原理图_10为例，系统完成3页图纸处理，识别110个元件实例、命中11个组合规则，并从文本层提取图签、控制信号和17行元件表。",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_metric_cards(
        doc,
        [
            ("3页", "处理范围", "图纸全页解析"),
            ("110个", "元件实例", "开放识别计数"),
            ("11组", "组合关系", "三类规则命中"),
            ("17行", "元件表", "文本层结构化"),
        ],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    add_text(p, "本次汇报重点", size=11, bold=True, color=NAVY)
    add_bullets(
        doc,
        [
            "展示三项已完成功能：元件识别、组合识别、文本层解析/OCR能力基础。",
            "说明当前检测指标的口径、已验证范围和仍需补齐的精度评价。",
            "提出面向历史图纸、元件知识和典型回路的混合检索建设规划。",
        ],
    )

    # Page 2 — capability overview
    doc.add_page_break()
    add_page_title(doc, "01 / 当前能力", "从“识别元件”走向“理解图纸”", "当前版本已具备三类结构化能力，并为后续检索形成数据底座。")
    flow = doc.add_table(rows=1, cols=5)
    flow.autofit = False
    flow_data = [
        ("图纸解析", "PDF分页、图像切片、文本层读取"),
        ("元件识别", "开放识别、同类归并、知识库名称修正"),
        ("组合识别", "按代号及功能关系匹配典型组合"),
        ("文本结构化", "图签、控制信号、元件表提取"),
        ("检索应用", "历史图纸、元件、回路和证据定位"),
    ]
    for idx, (name, detail) in enumerate(flow_data):
        cell = flow.cell(0, idx)
        set_cell_width(cell, 1872)
        set_cell_fill(cell, PALE_BLUE if idx < 4 else PALE_ORANGE)
        set_cell_margins(cell, top=160, bottom=160, start=90, end=90)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 4, "color": "B7C9E2"},
            bottom={"val": "single", "sz": 4, "color": "B7C9E2"},
            left={"val": "single", "sz": 4, "color": "B7C9E2"},
            right={"val": "single", "sz": 4, "color": "B7C9E2"},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, name, size=10, bold=True, color=NAVY if idx < 4 else ORANGE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p2, detail, size=7.5, color=DARK)

    doc.add_heading("已完成内容", level=2)
    completed = doc.add_table(rows=4, cols=4)
    completed_data = [
        ("模块", "已实现能力", "输出结果", "阶段状态"),
        ("元件识别", "按图块开放识别；按名称/代号归并数量；知识库RAG修正类别名称", "元件名称、代号、页码、数量、位置、置信度", "已实现"),
        ("组合识别", "基于代号一致性、元件类型及功能角色识别组合", "组合类型、组合代号、成员、证据、规则置信度", "已实现"),
        ("文本层解析", "读取PDF文本层，抽取图签、控制信号、控制方式及元件表", "结构化字段与表格数据", "已实现"),
    ]
    for r, row in enumerate(completed_data):
        for c, value in enumerate(row):
            completed.cell(r, c).text = value
    style_table(completed, widths=[1350, 3650, 2860, 1500], header=True, font_size=8.5)
    for r in range(1, 4):
        set_cell_fill(completed.cell(r, 3), PALE_GREEN)
        for p in completed.cell(r, 3).paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("关键优化思路", level=2)
    add_bullets(
        doc,
        [
            "开放识别阶段先完成同类元件归并与数量统计，例如“[代号A]：[临时类别A]：[10]”。",
            "知识库修正面向“类别”执行，将临时类别A统一修正为“指示灯”等标准名称，不再对10个实例逐一重复比对。",
            "组合识别消费结构化元件结果，避免再次调用视觉模型，降低大图纸处理时延与超时风险。",
        ],
    )
    add_callout(
        doc,
        "价值定位",
        "当前成果不是单一的图标分类器，而是一套可持续沉淀“对象、关系、文本、证据位置”的图纸理解流水线。",
    )

    # Page 3 — case overview
    doc.add_page_break()
    add_page_title(doc, "02 / 案例展示", "原理图_10：专用排烟风机控制原理", "案例包含主回路、控制回路、信号输入、状态指示和元件清单，适合验证多能力协同。")
    add_image(doc, assets["full"], 6.45, "图1  原理图_10第1页：主回路与控制回路（原始图纸）")
    case_table = doc.add_table(rows=3, cols=4)
    case_data = [
        ("图纸页数", "3页", "处理耗时", "146.66秒"),
        ("识别模式", "混合 / 视觉优先", "视觉模型请求", "7次"),
        ("知识库规模", "157个元件条目", "开放识别视图", "12/12成功"),
    ]
    for r, row in enumerate(case_data):
        for c, value in enumerate(row):
            case_table.cell(r, c).text = value
    style_table(case_table, widths=[1500, 2100, 1600, 4160], header=False, font_size=9)
    for r in range(3):
        for c in (0, 2):
            set_cell_fill(case_table.cell(r, c), LIGHT)
            for run in case_table.cell(r, c).paragraphs[0].runs:
                run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    add_text(
        p,
        "说明：处理耗时为该次历史任务记录，受图纸复杂度、服务并发、模型响应和缓存状态影响；本次数据用于展示工程运行情况，不作为性能SLA。",
        size=8.5,
        color=MID,
    )

    # Page 4 — component recognition
    doc.add_page_break()
    add_page_title(doc, "03 / 元件识别", "开放识别、数量归并与知识库修正", "系统先识别图纸中的视觉对象，再按类别和代号聚合，最后统一规范名称。")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_text(
        p,
        "在图1左侧主回路中，系统可识别QF、KM、F1、P1等对象；在右侧控制回路中，可继续识别按钮、继电器线圈/触点及状态指示灯。",
        size=9.5,
        color=DARK,
    )
    add_metric_cards(
        doc,
        [
            ("92条", "开放识别记录", "切片结果去重前"),
            ("23类", "开放类别", "类别级归并"),
            ("110个", "元件实例", "数量汇总"),
            ("82个", "修正实例", "14类RAG修正"),
        ],
    )
    doc.add_heading("原理图_10典型输出", level=2)
    examples = doc.add_table(rows=7, cols=4)
    example_data = [
        ("代号", "标准名称", "数量/关系", "识别用途"),
        ("QF", "三极断路器", "1", "主回路开关与保护"),
        ("KM", "接触器及线圈/触点", "跨页关联", "主回路驱动与控制联动"),
        ("F1", "热继电器/保护元件", "主元件+保护触点", "过载保护"),
        ("Y1、G1、R1", "指示灯", "3", "故障、停机、开机状态"),
        ("SF1、SS1", "启动/停止按钮", "各1", "启停控制输入"),
        ("X01～X04、XT", "接线端子", "多实例", "外部接线和信号接口"),
    ]
    for r, row in enumerate(example_data):
        for c, value in enumerate(row):
            examples.cell(r, c).text = value
    style_table(examples, widths=[1800, 2760, 1800, 3000], header=True, font_size=8.5)
    add_callout(
        doc,
        "效率优化",
        "RAG修正已按“每批12个类别、每类最多8个候选”执行。这里的批次对象应理解为归并后的类别/元件组，而不是对同类的每个物理实例重复检索。",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    # Page 5 — combination recognition
    doc.add_page_break()
    add_page_title(doc, "04 / 组合识别", "从单个元件到功能回路关系", "组合规则利用代号、元件角色和共同出现关系，输出可解释的功能组合。")
    add_image(doc, assets["control"], 3.35, "图2  控制回路局部：启停、继电器触点、线圈与状态指示")
    combo = doc.add_table(rows=4, cols=4)
    combo_data = [
        ("规则场景", "规则判断要点", "原理图_10命中", "规则置信度"),
        ("继电器/接触器线圈—触点组合", "相同代号同时出现线圈和辅助触点", "9组：1K1、1K2、2K1、2K2、K11、K12、K13、KA、KM", "98%"),
        ("电动机启动与保护组合", "断路器+接触器+热保护+负载共同出现", "1组：QF、KM、F1、P1", "96%"),
        ("启停控制及状态指示组合", "启动/停止、控制继电器和红黄绿指示共同构成控制链", "1组：SF1、SS1、继电器组、Y1/G1/R1", "94%"),
    ]
    for r, row in enumerate(combo_data):
        for c, value in enumerate(row):
            combo.cell(r, c).text = value
    style_table(combo, widths=[2200, 3100, 2900, 1160], header=True, font_size=8)
    for r in range(1, 4):
        set_cell_fill(combo.cell(r, 3), PALE_GREEN)
        combo.cell(r, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_callout(
        doc,
        "指标口径",
        "98%、96%、94%为规则内部的匹配强度/置信度，用于排序和人工复核提示；在建立人工标注集之前，不应等同于准确率、召回率或F1。",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    # Page 6 — text layer / OCR
    doc.add_page_break()
    add_page_title(doc, "05 / 文本层解析", "让图纸中的文字进入结构化数据", "当前优先读取矢量PDF文本层；图签、控制信号和元件表可直接用于筛选、核验与检索。")
    two_col = doc.add_table(rows=1, cols=2)
    two_col.autofit = False
    set_cell_width(two_col.cell(0, 0), 4680)
    set_cell_width(two_col.cell(0, 1), 4680)
    for cell in two_col.rows[0].cells:
        set_cell_margins(cell, top=50, bottom=50, start=60, end=60)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    p = two_col.cell(0, 0).paragraphs[0]
    add_text(p, "图签与控制信息", size=11, bold=True, color=NAVY)
    for text in [
        "图纸名称：专用排烟风机控制原理",
        "合同号：A17387",
        "原理图号：CDDT-6-DZ-.10",
        "页数：共3页，第1页",
        "信号：公共端、自动、电源、风机运行、风机故障",
        "控制方式：环控、就地、FAS启停、IBP盘启停、联锁风阀",
    ]:
        pp = two_col.cell(0, 0).add_paragraph(style="List Bullet")
        pp.paragraph_format.space_after = Pt(3)
        add_text(pp, text, size=8.8, color=DARK)
    p = two_col.cell(0, 1).paragraphs[0]
    add_text(p, "元件表提取", size=11, bold=True, color=NAVY)
    for text in [
        "结构化提取17行核心元件记录",
        "字段包括：序号、代号、元件名称、规格型号、数量、备注",
        "可与视觉识别结果交叉核验代号和数量",
        "典型记录：SS1停止按钮、SF1启动按钮、Y1/G1/R1指示灯",
    ]:
        pp = two_col.cell(0, 1).add_paragraph(style="List Bullet")
        pp.paragraph_format.space_after = Pt(3)
        add_text(pp, text, size=8.8, color=DARK)
    add_image(doc, assets["title"], 6.35, "图3  图签文本层提取区域")
    sample = doc.add_table(rows=5, cols=5)
    sample_data = [
        ("代号", "元件名称", "数量", "备注", "可用于检索的字段"),
        ("SS1", "按钮", "1", "停止", "停止按钮、SS1"),
        ("SF1", "按钮", "1", "启动", "启动按钮、SF1"),
        ("Y1", "指示灯", "1", "故障指示", "故障灯、黄色指示"),
        ("G1 / R1", "指示灯", "各1", "停机 / 开机", "运行状态、颜色、代号"),
    ]
    for r, row in enumerate(sample_data):
        for c, value in enumerate(row):
            sample.cell(r, c).text = value
    style_table(sample, widths=[1400, 1700, 1000, 2200, 3060], header=True, font_size=8.2)
    p = doc.add_paragraph()
    add_text(
        p,
        "边界：当前能力以带文本层的PDF为主。扫描件、照片型图纸和严重倾斜/低清晰度文件仍需接入版面分析与OCR引擎，并单独评估字符准确率。",
        size=8.5,
        color=ORANGE,
    )

    # Page 7 — metrics
    doc.add_page_break()
    add_page_title(doc, "06 / 阶段指标", "当前可报告指标与评价边界", "本阶段先证明流程可运行、结果可解释、模块可回归；下一阶段补齐标注集与正式精度指标。")
    indicators = doc.add_table(rows=10, cols=4)
    indicator_data = [
        ("指标维度", "原理图_10结果", "指标性质", "说明"),
        ("图纸处理", "3页全部完成", "覆盖指标", "主回路、控制回路、元件表"),
        ("开放识别视图", "12成功 / 0失败", "稳定性指标", "切片调用未出现失败"),
        ("元件实例", "110个", "规模指标", "开放识别汇总数量"),
        ("元件类别", "23类", "归并指标", "同类元件合并后类别数"),
        ("结构化结果", "24条", "输出指标", "按页与标准名称形成结果记录"),
        ("RAG名称修正", "14类 / 82个实例", "处理量指标", "类别级知识库规范化"),
        ("组合识别", "11组", "规则命中指标", "9+1+1三类组合"),
        ("文本元件表", "17行", "提取覆盖指标", "核心元件清单结构化"),
        ("自动化测试", "21项通过", "工程质量指标", "当前测试集回归通过率100%"),
    ]
    for r, row in enumerate(indicator_data):
        for c, value in enumerate(row):
            indicators.cell(r, c).text = value
    style_table(indicators, widths=[1900, 2200, 1800, 3460], header=True, font_size=8)
    doc.add_heading("尚不能直接下结论的指标", level=2)
    add_bullets(
        doc,
        [
            "元件识别准确率、召回率和F1：需要按元件实例建立人工真值标注。",
            "组合识别准确率和召回率：需要由电气专业人员确认真实组合边界和漏检情况。",
            "OCR字符准确率（CER/WER）：需要覆盖矢量PDF、扫描PDF、照片和低清图纸。",
            "端到端时延SLA：需要在固定硬件、固定模型、不同页数与复杂度分层压测。",
        ],
    )
    add_callout(
        doc,
        "建议的下一轮验收口径",
        "建立不少于50套、覆盖不同图纸类型与复杂度的标注集；按元件、组合、文本和检索四条线分别计算Precision、Recall、F1、CER/WER、Recall@K和MRR。",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    # Page 8 — retrieval plan
    doc.add_page_break()
    add_page_title(doc, "07 / 后续规划", "建设可追溯的图纸与知识检索能力", "目标不是只返回答案，而是返回“图纸—页码—区域—元件/组合—文本证据”的完整定位链。")
    roadmap = doc.add_table(rows=5, cols=4)
    roadmap_data = [
        ("建设层", "主要内容", "用户可获得的能力", "建议优先级"),
        ("数据规范层", "统一图纸、页面、元件实例、组合关系、图签和元件表的数据模型", "跨任务统一查询、结果可复用", "P0"),
        ("混合索引层", "代号/型号精确检索 + 文本BM25 + 图像/文本向量 + 关系索引", "既能查“QF”，也能查“排烟风机启停回路”", "P0"),
        ("检索与重排层", "多路召回、规则过滤、语义重排、证据聚合", "提高历史相似图和典型回路命中率", "P1"),
        ("应用与闭环层", "图纸问答、相似图纸、元件替代、组合模板、人工反馈", "面向设计复用、审查和运维定位", "P1/P2"),
    ]
    for r, row in enumerate(roadmap_data):
        for c, value in enumerate(row):
            roadmap.cell(r, c).text = value
    style_table(roadmap, widths=[1600, 3450, 3250, 1060], header=True, font_size=8.2)
    for r in range(1, 5):
        roadmap.cell(r, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_fill(roadmap.cell(r, 3), PALE_BLUE if r < 3 else LIGHT)

    doc.add_heading("建议分阶段推进", level=2)
    phases = doc.add_table(rows=3, cols=3)
    phase_data = [
        ("阶段一：可查", "历史结果规范化；按图号、名称、代号、型号、页码和文本字段检索；结果跳转证据区域", "形成基础检索门户与评价集"),
        ("阶段二：相似", "引入多模态向量和重排；支持相似元件、相似图纸、相似组合回路检索", "提升设计复用和问题定位效率"),
        ("阶段三：可问", "关系图谱与检索增强生成；支持自然语言问答、差异对比和审查提示", "形成面向业务的智能图纸助手"),
    ]
    for r, row in enumerate(phase_data):
        for c, value in enumerate(row):
            phases.cell(r, c).text = value
    style_table(phases, widths=[1900, 4760, 2700], header=False, font_size=8.5)
    for r in range(3):
        set_cell_fill(phases.cell(r, 0), PALE_BLUE)
        for run in phases.cell(r, 0).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)

    doc.add_heading("建议管理层关注的三项决策", level=2)
    add_bullets(
        doc,
        [
            "明确首批检索场景优先级：历史图纸定位、元件知识查询、典型组合回路复用。",
            "安排电气专业人员参与标注与抽检，建立统一的准确率验收口径。",
            "确定历史图纸数据范围、权限和脱敏规则，保证检索结果可用且可控。",
        ],
    )
    add_callout(
        doc,
        "近期里程碑建议",
        "以现有result历史任务为首批数据源，先完成结构化入库和精确检索；同步建立原理图_05、_07、_10等样例的人工标注基线，再逐步增加语义与多模态检索。",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_report()
